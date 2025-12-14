from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import RGBColor

def create_word_template(filename="report_template.docx"):
    doc = Document()
    
    # helper for formatting
    def add_field(paragraph, text, bold=False, color=None, size=None):
        run = paragraph.add_run(text)
        if bold: run.bold = True
        if color: run.font.color.rgb = color
        if size: run.font.size = size
        return run

    # 1. Header (Blue Banner Simulation)
    # Word doesn't do "banners" easily without tables or header editing.
    # We'll use a simple text header for now.
    
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = add_field(header_para, "NATIONAL AIR QUALITY MONITORING BUREAU", bold=True, size=Pt(16))
    run.font.color.rgb = RGBColor(26, 35, 126) # #1a237e
    
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(sub_para, "MINISTRY OF ENVIRONMENT & FORESTS | GOVT. OF INDIA", size=Pt(10))
    
    doc.add_paragraph() # Spacer
    
    # 2. Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(title, "REGULATORY COMPLIANCE REPORT", bold=True, size=Pt(20))
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(subtitle, "JURISDICTION: {{ city }}", size=Pt(12))
    
    doc.add_paragraph() # Spacer
    
    # 3. Metadata Table
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    cells = table.rows[0].cells
    cells[0].text = 'Report Number:'
    cells[1].text = 'AQI-{{ timestamp }}-{{ city[:3] }}'
    
    cells = table.rows[1].cells
    cells[0].text = 'Date of Issue:'
    cells[1].text = '{{ date }}'
    
    cells = table.rows[2].cells
    cells[0].text = 'Monitoring Period:'
    cells[1].text = '{{ start_date }} to {{ end_date }}'
    
    cells = table.rows[3].cells
    cells[0].text = 'Overall Status:'
    cells[1].text = '{{ category }}'
    
    doc.add_paragraph()
    
    # 4. Metrics Highlight
    p = doc.add_paragraph()
    add_field(p, "KEY METRICS", bold=True, size=Pt(14))
    
    # Using a table for dashboard-like layout
    d_table = doc.add_table(rows=2, cols=3)
    d_table.style = 'Table Grid'
    
    # Headers
    d_table.rows[0].cells[0].text = "AQI"
    d_table.rows[0].cells[1].text = "Cigarettes/Day"
    d_table.rows[0].cells[2].text = "Main Pollutant"
    
    # Values (Placeholders)
    d_table.rows[1].cells[0].text = "{{ aqi }}"
    d_table.rows[1].cells[1].text = "{{ cigarettes }}"
    d_table.rows[1].cells[2].text = "{{ dominant_pollutant }}"
    
    doc.add_paragraph()
    
    # 5. Narrative (Rich Text tags)
    doc.add_heading('Executive Summary', level=2)
    doc.add_paragraph("{{ summary_text }}")
    
    doc.add_heading('Policy Recommendations', level=2)
    doc.add_paragraph("{{ recommendations }}") # docxtpl can interpret newlines? better use r_recommendations object
    
    # 6. Compliance Data (Loop example)
    doc.add_heading('Compliance Data', level=2)
    
    # Table with loop
    c_table = doc.add_table(rows=1, cols=4)
    c_table.style = 'Table Grid'
    hdr_cells = c_table.rows[0].cells
    hdr_cells[0].text = 'Pollutant'
    hdr_cells[1].text = 'Standard'
    hdr_cells[2].text = 'Observed'
    hdr_cells[3].text = 'Status'
    
    # Template Row for Jinja loop
    # In docxtpl, you normally put `{% for p in pollutants %}` inside the first cell
    # and `{% endfor %}` in the last cell (or use specific row tags).
    # Since writing complex tags via python-docx is error-prone, 
    # we'll just put simple placeholders for now and let the user customize.
    # Simpler: Just a static table for NO2, SO2 etc since we know them.
    
    pollutants = ['NO2', 'SO2', 'CO', 'O3', 'PM2.5', 'PM10']
    for pol in pollutants:
        row = c_table.add_row().cells
        safe_pol = pol.replace('.', '_')
        row[0].text = pol
        row[1].text = f"{{{{ {safe_pol}_std }}}}"
        row[2].text = f"{{{{ {safe_pol}_val }}}}"
        row[3].text = f"{{{{ {safe_pol}_status }}}}"

    doc.add_paragraph()
    
    # 7. Signature
    doc.add_paragraph()
    sig_p = doc.add_paragraph()
    sig_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_field(sig_p, "Authorized Signatory", bold=True)
    doc.add_paragraph("_" * 30).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph("Chief Toxicologist").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.save(filename)
    print(f"Created Word Template: {filename}")

if __name__ == "__main__":
    create_word_template()
